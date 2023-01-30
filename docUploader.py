#import libs
import openai 
import streamlit as st
import io
import docx
import pandas as pd
from transformers import GPT2TokenizerFast\


# pip install streamlit-chat  
from streamlit_chat import message

openai.api_key = st.secrets['openai-secret']

#add title to 
st.title("chatBot : OpenAI")

#Sidebar
with st.sidebar:
    # Create a slider to control the temperature of the model
    temperature = st.slider("Increase the temperature for wider variations", 
        min_value=0.0, 
        max_value=1.0, 
        value=0.8, 
        step=0.1)
    max_tokens = st.number_input('Insert max tokens (1 token = 3/4 words)',
        min_value=0, 
        max_value=3000,
        value=2000)

#Initialising session
if 'generated' not in st.session_state:
    st.session_state['generated'] = []

if 'past' not in st.session_state:
    st.session_state['past'] = []

if 'history' not in st.session_state:
    st.session_state['history'] = []

if 'text' not in st.session_state:
    st.session_state.text = ""

def generate_response(prompt):
        completions = openai.Completion.create(
            engine="text-davinci-003", #Select Model
            prompt = prompt,
            temperature = temperature, #0-1 Higher the temp, the more liberties it takes 
            top_p = 1,
            max_tokens = max_tokens, #1 Token = 3/4 word
            frequency_penalty = 0,
            presence_penalty = 0
    )
        return completions.choices[0].text

#Form for user input
def update():
    st.session_state.text += st.session_state.text_value

with st.form(key='my_form',clear_on_submit=True):
    uploaded_file = st.file_uploader("Choose a doc file")
    if uploaded_file is not None:
        doc = docx.Document(uploaded_file)
        paras = '\n'.join([p.text for p in doc.paragraphs if p.text]   )
        paras.splitlines()
        st.write(paras)
    st.text_input('Enter your prompt and click on submit', value="", key='text_value')
    submit = st.form_submit_button(label='Submit', on_click=update)

#Get user response
if uploaded_file is not None and st.session_state.text != "":
    #Get user response
    st.session_state.text = st.session_state.text + paras
    uploaded_file = None

user_input = st.session_state.text

st.session_state.text = ""

if user_input:
        
    #Store the input
    st.session_state.history.append(user_input)

    #Get the conversation history
    conversation_history = '\n'.join(st.session_state['history'])
    output = generate_response(conversation_history)

    #Store the chat
    st.session_state.past.append(user_input)
    st.session_state.generated.append(output)
    st.session_state.history.append(output)

# Create an instance of a word document
def list_to_word_doc(items, doc_name):
    doc = docx.Document()
    for item in items:
        doc.add_paragraph(item)
    return doc

#Chat history downloader
doc_download = list_to_word_doc(st.session_state['history'], 'chat_history')

bio = io.BytesIO()
doc_download.save(bio)
if doc_download:
    st.download_button(
        label="Download Chat History",
        data=bio.getvalue(),
        file_name="chat_history.docx",
        mime="docx"
    )

#display no of tokens
tokenizer = GPT2TokenizerFast.from_pretrained("gpt2")
number_of_tokens = len(tokenizer(''.join(st.session_state['history']))['input_ids'])

st.text('Number of tokens left: '+ str(max_tokens - number_of_tokens))

if st.session_state['generated']:
    
    for i in range(len(st.session_state['generated'])-1, -1, -1):
        message(st.session_state["generated"][i], key=str(i))
        message(st.session_state['past'][i], is_user=True, key=str(i) + '_user')
